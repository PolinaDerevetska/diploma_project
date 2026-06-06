
import re
import logging
from docx import Document

import database as db

log = logging.getLogger(__name__)

CATEGORY_MAP = {
    "B1.1": "B1.1",
    "В1.1": "B1.1",
    "B1.3": "B1.3",
    "В1.3": "B1.3",
    "B2":   "B2",
    "В2":   "B2",

    "B1.1": "B1.1",
    "B1.3": "B1.3",
    "B2":   "B2",
}

def _parse_module_from_paragraph(text: str) -> tuple[str, str] | None:
    m = re.match(r'[Мм]одуль\s+(\d+)\s*[-–.]\s*(.*)', text.strip())
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None

def _parse_section_from_row(cells: list[str]) -> tuple[str, str] | None:
    if len(cells) < 2:
        return None
    raw = cells[1].strip()
    if not raw:
        return None

    m = re.match(r'^([\d]+\.[\d.]+)\s*(.*)', raw)
    if not m:
        return None

    base_code = m.group(1).strip()
    rest      = m.group(2).strip()
    name      = rest if rest else base_code

    suffix = re.search(r'\(([a-zA-Zа-яА-ЯіІїЇєЄёЁ])\)\s*$', name)
    if suffix:
        code = f"{base_code}({_normalize_suffix(suffix.group(1))})"
    else:
        code = base_code

    return code, name

_CYR_SUFFIX_MAP = {"а": "a", "б": "b", "в": "c", "г": "d",
                   "с": "c", "е": "e", "d": "d"}

def _normalize_suffix(ch: str) -> str:
    ch = ch.lower()
    return _CYR_SUFFIX_MAP.get(ch, ch)

def load_quotas_from_docx(filepath: str, db_path: str = db.DB_PATH) -> None:
    log.info("Завантаження норм із: %s", filepath)
    doc = Document(filepath)

    cat_ids: dict[str, int] = {}
    for cat_code in ("B1.1", "B1.3", "B2"):
        cat_ids[cat_code] = db.insert_category(cat_code, db_path)

    module_code: str | None = None
    module_name: str | None = None
    module_id:   int | None = None

    table_idx = 0
    tables = doc.tables

    all_paragraphs = doc.paragraphs
    para_idx = 0

    for para in all_paragraphs:
        text = para.text.strip()
        if not text:
            continue
        parsed = _parse_module_from_paragraph(text)
        if parsed:
            module_code, module_name = parsed
            log.debug("Модуль: %s — %s", module_code, module_name)
            module_id = db.insert_module(module_code, module_name, db_path)

    _load_quotas_from_tables(doc, cat_ids, db_path)

    log.info("Норми завантажено успішно")

def _load_quotas_from_tables(doc: Document, cat_ids: dict[str, int], db_path: str) -> None:

    from docx.oxml.ns import qn as _qn
    body = doc.element.body

    current_module_id: int | None = None
    current_module_code: str | None = None
    table_index = 0

    tables_list = doc.tables

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":

            text = "".join(r.text or "" for r in child.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
            text = text.strip()
            parsed = _parse_module_from_paragraph(text)
            if parsed:
                code, name = parsed
                current_module_id   = db.insert_module(code, name, db_path)
                current_module_code = code
                log.debug("Поточний модуль: %s", name)

        elif tag == "tbl":

            if table_index >= len(tables_list):
                break
            table = tables_list[table_index]
            table_index += 1

            if current_module_id is None:
                log.warning("Таблиця %d без модуля, пропускаємо", table_index)
                continue

            _process_quota_table(table, current_module_id, cat_ids, db_path)

def _process_quota_table(table, module_id: int, cat_ids: dict[str, int], db_path: str) -> None:
    if not table.rows:
        return

    col_cat_map: dict[int, str] = {}
    header_found = False

    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]

        if any(c in ("В1.1", "B1.1", "В1.3", "B1.3", "B2", "В2", "B1.1", "B1.3", "B2") for c in cells):
            for i, cell_text in enumerate(cells):
                norm = cell_text.replace("В", "B").strip()
                if norm in CATEGORY_MAP:
                    col_cat_map[i] = CATEGORY_MAP[norm]
            header_found = True
            continue

        if not header_found:
            continue

        if not cells[0].isdigit():
            continue

        parsed = _parse_section_from_row(cells)
        if not parsed:
            continue
        sec_code, sec_name = parsed

        section_id = db.insert_section(module_id, sec_code, sec_name, db_path)

        for col_idx, cat_code in col_cat_map.items():
            if col_idx >= len(cells):
                continue
            count_str = cells[col_idx].strip().replace("−", "-").replace("–", "0")
            if count_str in ("-", ""):
                count = 0
            else:
                try:
                    count = int(count_str)
                except ValueError:
                    count = 0

            cat_id = cat_ids.get(cat_code)
            if cat_id is not None and count > 0:
                db.insert_quota(section_id, cat_id, count, db_path)

    log.debug("  Оброблено таблицю для модуля id=%d", module_id)

if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) < 2:
        print("Використання: python quota_loader.py <шлях_до_файлу>")
        sys.exit(1)
    db.init_db()
    load_quotas_from_docx(sys.argv[1])
    print(db.get_db_stats())
