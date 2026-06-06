
import os
import io
import re
import struct
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

from generator import TestVariant, Question, QuestionImage

def _unique_filepath(output_dir: str, base_name: str) -> str:
    candidate = os.path.join(output_dir, f"{base_name}.docx")
    if not os.path.exists(candidate):
        return candidate
    n = 2
    while True:
        candidate = os.path.join(output_dir, f"{base_name}_{n:03d}.docx")
        if not os.path.exists(candidate):
            return candidate
        n += 1

_emf_counter = 0
_WMF_MAGIC = 0x9AC6CDD7

def _get_emf_dimensions_emu(data: bytes) -> tuple[int, int]:
    try:
        if len(data) >= 22 and struct.unpack_from("<I", data, 0)[0] == _WMF_MAGIC:
            left, top, right, bottom = struct.unpack_from("<hhhh", data, 6)
            inch = struct.unpack_from("<H", data, 14)[0] or 96
            w = (right - left) * 914_400 // inch
            h = (bottom - top) * 914_400 // inch
            if w > 0 and h > 0:
                return w, h
        left, top, right, bottom = struct.unpack_from("<iiii", data, 24)
        w = (right - left) * 360
        h = (bottom - top) * 360
        if w > 0 and h > 0:
            return w, h
    except Exception:
        pass
    return 2_880_000, 2_160_000

def _insert_emf_picture(paragraph, emf_data: bytes, doc_part, max_width_emu: int = 2_880_000):
    global _emf_counter
    _emf_counter += 1
    width_emu, height_emu = _get_emf_dimensions_emu(emf_data)
    if width_emu > max_width_emu:
        height_emu = int(height_emu * max_width_emu / width_emu)
        width_emu = max_width_emu
    uri = PackURI(f"/word/media/emf_{_emf_counter}.emf")
    emf_part = Part(uri, "image/x-emf", emf_data, doc_part.package)
    rId = doc_part.relate_to(
        emf_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
    )
    xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{_emf_counter}" name="emf{_emf_counter}"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="emf{_emf_counter}"/>'
        f'<pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing>'
    )
    run = paragraph.add_run()
    run._element.append(etree.fromstring(xml.encode("utf-8")))

def _set_font(run, size_pt: float, bold: bool = False, italic: bool = False,
              color: tuple | None = None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor(*color)

def _add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, 12)
    return p

def _set_margins(doc, left_cm=3.0, right_cm=2.0, top_cm=2.0, bottom_cm=2.0):
    for section in doc.sections:
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)

def _cell_font(cell, text, size_pt=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_font(run, size_pt, bold=bold)

def _set_cell_shading(cell, fill_hex: str = "D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def _set_cell_margins(cell, top=20, bottom=20, left=40, right=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

_IMG_MARKER_RE = re.compile(r'\[IMG_(\d+)\]')

def _insert_single_image(paragraph, img, doc_part, is_option):
    max_w = 1_440_000 if is_option else 2_880_000
    if img.ext in ("emf", "wmf"):
        _insert_emf_picture(paragraph, img.data, doc_part, max_width_emu=max_w)
    else:
        run = paragraph.add_run()
        try:
            width = Inches(1.5) if is_option else Inches(3.5)
            run.add_picture(io.BytesIO(img.data), width=width)
        except Exception:
            _set_font(run, 10, italic=True)
            run.text = "[зображення]"

def _add_text_with_images(paragraph, text, images, context, doc_part,
                          font_size=12, bold=False):
    is_option = context in ("option_a", "option_b", "option_c")
    ctx_imgs = {img.img_index: img for img in images if img.context == context}
    sorted_imgs = sorted(ctx_imgs.values(), key=lambda i: i.img_index)
    if _IMG_MARKER_RE.search(text):
        img_iter = iter(sorted_imgs)
        parts = _IMG_MARKER_RE.split(text)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    run = paragraph.add_run(part)
                    _set_font(run, font_size, bold=bold)
            else:
                img = next(img_iter, None)
                if img:
                    _insert_single_image(paragraph, img, doc_part, is_option)
    else:
        if text:
            run = paragraph.add_run(text)
            _set_font(run, font_size, bold=bold)
        for img in sorted_imgs:
            _insert_single_image(paragraph, img, doc_part, is_option)

def _add_title_block(doc, variant, is_answer_key=False):

    smy_code = "40/03-11-14" if is_answer_key else "40/03-11-07"
    p_smy = doc.add_paragraph()
    p_smy.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_smy.paragraph_format.space_before = Pt(0)
    p_smy.paragraph_format.space_after = Pt(0)
    r_smy = p_smy.add_run(f"СМЯ КРФК КАІ {smy_code}")
    _set_font(r_smy, 11, bold=True, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    if is_answer_key:
        r_t1 = p_title.add_run("Вірні відповіді на тестове екзаменаційне завдання № ")
        _set_font(r_t1, 14, bold=True)
        r_t2 = p_title.add_run("______________")
        _set_font(r_t2, 14, bold=True, italic=True)
        r_t2.font.underline = True
    else:
        r_t1 = p_title.add_run("Екзаменаційне завдання № ")
        _set_font(r_t1, 14, bold=True)
        r_t2 = p_title.add_run("______________")
        _set_font(r_t2, 14, bold=True, italic=True)
        r_t2.font.underline = True

    p_cat = doc.add_paragraph()
    p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cat.paragraph_format.space_before = Pt(0)
    p_cat.paragraph_format.space_after = Pt(6)
    r_cat = p_cat.add_run(f"Категорія: {variant.category_code}   |   {variant.student_name}")
    _set_font(r_cat, 12, bold=True)

def _add_question_student(doc, num, q):

    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr.paragraph_format.space_before = Pt(6)
    p_hdr.paragraph_format.space_after = Pt(0)
    r_hdr = p_hdr.add_run(f"Завдання {num}")
    _set_font(r_hdr, 12, bold=True)
    r_hdr.font.underline = True

    p_q = doc.add_paragraph()
    p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_q.paragraph_format.space_before = Pt(0)
    p_q.paragraph_format.space_after = Pt(0)
    p_q.paragraph_format.left_indent = Cm(1)
    _add_text_with_images(p_q, q.question_text, q.images, "question", doc.part,
                          font_size=12, bold=False)

    for label, text, ctx in [
        ("А)", q.option_a or "", "option_a"),
        ("В)", q.option_b or "", "option_b"),
        ("С)", q.option_c or "", "option_c"),
    ]:
        op = doc.add_paragraph()
        op.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        op.paragraph_format.left_indent = Cm(1)
        op.paragraph_format.space_before = Pt(0)
        op.paragraph_format.space_after = Pt(0)
        rl = op.add_run(f"{label} ")
        _set_font(rl, 12, bold=False)
        _add_text_with_images(op, text, q.images, ctx, doc.part,
                              font_size=12, bold=False)

_BLOCK_SIZE = 45
_BLOCKS_PER_ROW = 4

def _add_answer_sheet(doc, variant):
    doc.add_page_break()

    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ph.paragraph_format.space_before = Pt(0)
    ph.paragraph_format.space_after = Pt(0)
    rh = ph.add_run("СМЯ КРФК КАІ 40/03-11-13")
    _set_font(rh, 11, bold=True, italic=True)

    ph2 = doc.add_paragraph()
    ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph2.paragraph_format.space_before = Pt(0)
    rh2 = ph2.add_run("Бланк відповідей на тестові питання")
    _set_font(rh2, 14, bold=True)
    ph2.paragraph_format.space_after = Pt(6)

    info_tbl = doc.add_table(rows=3, cols=1)
    info_tbl.style = "Table Grid"
    labels = [
        f"Назва курсу та категорія:  Part 147 — {variant.category_code}",
        "Прізвище ім'я по-батькові:  ",
        "Дата складання екзамену:  _____________________",
    ]
    for i, lbl in enumerate(labels):
        cell = info_tbl.rows[i].cells[0]
        _cell_font(cell, lbl, size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

    pn = doc.add_paragraph()
    rn = pn.add_run("Примітка. Закресліть хрестиком правильну відповідь.")
    _set_font(rn, 10, italic=True)
    pn.paragraph_format.space_before = Pt(6)
    pn.paragraph_format.space_after = Pt(6)

    n = variant.total_questions
    cols_per_block = 5
    total_cols = _BLOCKS_PER_ROW * cols_per_block

    col_w = [600, 400, 400, 400, 500] * _BLOCKS_PER_ROW

    rows_in_table = _BLOCK_SIZE + 1

    total_blocks = (n + _BLOCK_SIZE - 1) // _BLOCK_SIZE
    num_pages = (total_blocks + _BLOCKS_PER_ROW - 1) // _BLOCKS_PER_ROW

    for page_idx in range(num_pages):
        if page_idx > 0:
            doc.add_page_break()
            ph_c = doc.add_paragraph()
            ph_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rh_c = ph_c.add_run("Бланк відповідей (продовження)")
            _set_font(rh_c, 12, bold=True)

        tbl = doc.add_table(rows=rows_in_table, cols=total_cols)
        tbl.style = "Table Grid"

        hdr = tbl.rows[0]
        for b in range(_BLOCKS_PER_ROW):
            base = b * cols_per_block
            for ci, lbl in enumerate(["№", "А", "В", "С", "К"]):
                c = hdr.cells[base + ci]
                _cell_font(c, lbl, size_pt=9, bold=True)
                _set_cell_margins(c, top=20, bottom=20, left=20, right=20)
                _set_cell_shading(c, "D9D9D9")

                tc = c._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_w[base + ci]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

        for row_idx in range(1, rows_in_table):
            row = tbl.rows[row_idx]
            for b in range(_BLOCKS_PER_ROW):
                block_num = page_idx * _BLOCKS_PER_ROW + b
                q_num = block_num * _BLOCK_SIZE + row_idx
                base = b * cols_per_block
                c_num = row.cells[base]
                _set_cell_margins(c_num, top=15, bottom=15, left=20, right=20)
                if q_num <= n:
                    _cell_font(c_num, str(q_num), size_pt=9)

                    for ci, abc_lbl in enumerate(["А", "В", "С"], start=1):
                        c = row.cells[base + ci]
                        _cell_font(c, abc_lbl, size_pt=9)
                        _set_cell_margins(c, top=15, bottom=15, left=20, right=20)

                    c_k = row.cells[base + 4]
                    _cell_font(c_k, "", size_pt=9)
                    _set_cell_margins(c_k, top=15, bottom=15, left=20, right=20)
                else:
                    for ci in range(5):
                        c = row.cells[base + ci]
                        _set_cell_margins(c, top=15, bottom=15, left=20, right=20)

    doc.add_paragraph()
    footer_tbl = doc.add_table(rows=4, cols=1)
    footer_tbl.style = "Table Grid"
    footer_rows = [
        "Підпис особи, що здає екзамен: ___________",
        "Час початку екзамену: __________________",
        "Дата перевірки тестової роботи: _________",
        "Прізвище та підпис екзаменатора: ________",
    ]
    for i, lbl in enumerate(footer_rows):
        cell = footer_tbl.rows[i].cells[0]
        _cell_font(cell, lbl, size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

def _add_descriptive_sheet(doc, variant, desc_questions):
    if not desc_questions:
        return

    doc.add_page_break()

    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run("СМЯ КРФК КАІ 40/03-11-13")
    _set_font(rh, 9)

    ph2 = doc.add_paragraph()
    ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh2 = ph2.add_run("Бланк відповідей на описові питання")
    _set_font(rh2, 14, bold=True)
    ph2.paragraph_format.space_after = Pt(6)

    info_tbl = doc.add_table(rows=3, cols=1)
    info_tbl.style = "Table Grid"
    for i, lbl in enumerate([
        f"Назва курсу та категорія:  Part 147 — {variant.category_code}",
        "Прізвище та ім'я по батькові:  ",
        "Дата складання екзамену:  _____________________",
    ]):
        cell = info_tbl.rows[i].cells[0]
        _cell_font(cell, lbl, size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

    doc.add_paragraph()

    pq = doc.add_paragraph()
    rq = pq.add_run("Екзаменаційні запитання:")
    _set_font(rq, 12, bold=True)

    by_module: dict[str, list] = {}
    for q in desc_questions:
        by_module.setdefault(q.module_code, []).append(q)

    module_names = {
        '7': 'Модуль 7 — Практика технічного обслуговування',
        '9': 'Модуль 9 — Людський фактор',
        '10': 'Модуль 10 — Авіаційне законодавство',
    }

    num = 1
    for mod_code in sorted(by_module.keys()):
        questions = by_module[mod_code]
        pm = doc.add_paragraph()
        rm = pm.add_run(module_names.get(mod_code, f"Модуль {mod_code}"))
        _set_font(rm, 11, bold=True, italic=True)
        pm.paragraph_format.space_before = Pt(8)

        for q in questions:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            rnum = p.add_run(f"{num}. ")
            _set_font(rnum, 11, bold=True)
            _add_text_with_images(p, q.question_text, q.images, "question",
                                  doc.part, font_size=11)
            num += 1

    doc.add_paragraph()
    pa = doc.add_paragraph()
    ra = pa.add_run("Відповіді:")
    _set_font(ra, 12, bold=True)

    for i in range(1, num):
        pans = doc.add_paragraph()
        pans.paragraph_format.space_before = Pt(4)
        rans = pans.add_run(f"{i}. ")
        _set_font(rans, 11, bold=True)

        for _ in range(3):
            pline = doc.add_paragraph()
            pline.paragraph_format.space_before = Pt(0)
            pline.paragraph_format.space_after = Pt(0)
            rline = pline.add_run("_" * 80)
            _set_font(rline, 11)

    doc.add_paragraph()
    foot_tbl = doc.add_table(rows=5, cols=1)
    foot_tbl.style = "Table Grid"
    for i, lbl in enumerate([
        "Підпис особи, що здає екзамен: ___________",
        "Час початку екзамену: __________________",
        "Час закінчення екзамену: ________________",
        "Дата перевірки описової роботи: _________",
        "Відсоток вірних відповідей: _____________",
    ]):
        cell = foot_tbl.rows[i].cells[0]
        _cell_font(cell, lbl, size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

def export_student_docx(
    variant: TestVariant,
    output_dir: str,
    filename: str | None = None,
    descriptive: list | None = None,
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    safe_name = variant.student_name.replace(" ", "_").replace("/", "-")
    if filename is None:
        filename = f"Тест_{variant.category_code}_{safe_name}"
    filepath = _unique_filepath(output_dir, filename)

    doc = Document()
    _set_margins(doc)
    _add_title_block(doc, variant, is_answer_key=False)

    for i, q in enumerate(variant.questions, start=1):
        _add_question_student(doc, i, q)

    _add_answer_sheet(doc, variant)

    if descriptive:
        _add_descriptive_sheet(doc, variant, descriptive)

    doc.save(filepath)
    return filepath

def export_answer_key_docx(
    variant: TestVariant,
    output_dir: str,
    filename: str | None = None,
    descriptive: list | None = None,
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    safe_name = variant.student_name.replace(" ", "_").replace("/", "-")
    if filename is None:
        filename = f"Ключ_{variant.category_code}_{safe_name}"
    filepath = _unique_filepath(output_dir, filename)

    doc = Document()
    _set_margins(doc)
    _add_title_block(doc, variant, is_answer_key=True)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    hdr = tbl.rows[0].cells
    col_widths = [600, 6500, 900]
    for ci, (h, w) in enumerate(zip(["№", "Питання", "Відп."], col_widths)):
        c = hdr[ci]
        _cell_font(c, h, size_pt=10, bold=True)
        _set_cell_margins(c)
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    ans_labels = {"A": "А", "B": "В", "C": "С"}

    for i, q in enumerate(variant.questions, start=1):
        row = tbl.add_row()
        cells = row.cells

        _cell_font(cells[0], str(i), size_pt=9)
        _set_cell_margins(cells[0])

        q_text = q.question_text.replace('\n', ' ')
        if len(q_text) > 120:
            q_text = q_text[:117] + "..."
        p_q = cells[1].paragraphs[0]
        p_q.paragraph_format.space_before = Pt(0)
        p_q.paragraph_format.space_after = Pt(0)
        r_q = p_q.add_run(q_text)
        _set_font(r_q, 8)

        if q.images:
            r_img = p_q.add_run(" [IMG]")
            _set_font(r_img, 8, italic=True)
        _set_cell_margins(cells[1])

        ans = ans_labels.get(q.correct_answer, q.correct_answer)
        _cell_font(cells[2], ans, size_pt=10, bold=True)
        _set_cell_margins(cells[2])

    if descriptive:
        doc.add_page_break()
        pd = doc.add_paragraph()
        rd = pd.add_run("Ключ відповідей (описова частина — М7/М9/М10)")
        _set_font(rd, 13, bold=True)
        pd.paragraph_format.space_after = Pt(8)

        by_module: dict[str, list] = {}
        for q in descriptive:
            by_module.setdefault(q.module_code, []).append(q)

        module_names = {
            '7': 'Модуль 7 — Практика технічного обслуговування',
            '9': 'Модуль 9 — Людський фактор',
            '10': 'Модуль 10 — Авіаційне законодавство',
        }

        num = 1
        for mod_code in sorted(by_module.keys()):
            pm = doc.add_paragraph()
            rm = pm.add_run(module_names.get(mod_code, f"Модуль {mod_code}"))
            _set_font(rm, 11, bold=True, italic=True)
            pm.paragraph_format.space_before = Pt(8)

            for q in by_module[mod_code]:

                p_q = doc.add_paragraph()
                p_q.paragraph_format.space_before = Pt(4)
                rn = p_q.add_run(f"{num}. ")
                _set_font(rn, 10, bold=True)
                _add_text_with_images(p_q, q.question_text, q.images, "question",
                                      doc.part, font_size=10)

                ans_key = q.correct_answer.upper()
                opt_map = {'A': q.option_a, 'B': q.option_b, 'C': q.option_c}
                ans_text = opt_map.get(ans_key, '')
                ctx_map = {'A': 'option_a', 'B': 'option_b', 'C': 'option_c'}

                p_a = doc.add_paragraph()
                p_a.paragraph_format.left_indent = Cm(1)
                p_a.paragraph_format.space_before = Pt(2)
                p_a.paragraph_format.space_after = Pt(4)
                ra = p_a.add_run(f"Відповідь ({ans_labels.get(ans_key, ans_key)}): ")
                _set_font(ra, 10, bold=True, color=(0, 112, 0))
                _add_text_with_images(p_a, ans_text or '—', q.images,
                                      ctx_map.get(ans_key, 'question'),
                                      doc.part, font_size=10)
                num += 1

    doc.save(filepath)
    return filepath

if __name__ == "__main__":
    import sys
    sys.path.insert(0, os.path.dirname(__file__))
    from generator import generate_test
    v = generate_test("B1.1", "Тестовий Здобувач")
    out = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
    sf = export_student_docx(v, out)
    kf = export_answer_key_docx(v, out)
    print(f"Тест: {sf}")
    print(f"Ключ: {kf}")
