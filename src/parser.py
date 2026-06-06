
import os
import re
import logging
from docx import Document
from docx.oxml.ns import qn

import database as db

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)

ANSWER_MAP = {
    "А": "A", "A": "A",
    "В": "B", "B": "B", "Б": "B",
    "С": "C", "C": "C",
}

OPTION_KEYS = {
    "А": "option_a", "A": "option_a",
    "В": "option_b", "B": "option_b",
    "Б": "option_b",
    "С": "option_c", "C": "option_c",
}

_VML_NS = "urn:schemas-microsoft-com:vml"
_R_NS   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_EXT_MAP = {"jpeg": "jpg", "tiff": "tif", "x-emf": "emf",
            "x-wmf": "wmf", "x-msmetafile": "wmf"}

def _is_excluded_row(row) -> bool:
    full_text = " ".join(c.text for c in row.cells).lower()
    if "вилучено" in full_text:
        return True
    for cell in row.cells:
        tcPr = cell._tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            continue
        fill = (shd.get(qn("w:fill")) or "").upper()
        if fill == "FF0000":
            return True
    return False

def _extract_cell_content(cell, doc: Document) -> tuple[str, list[tuple[bytes, str]]]:
    para_texts: list[str] = []
    images: list[tuple[bytes, str]] = []

    for para in cell.paragraphs:
        parts: list[str] = []

        for child in para._element:
            local_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if local_tag != "r":
                continue

            for t_node in child.findall(f"{{{_W_NS}}}t"):
                parts.append(t_node.text or "")

            for blip in child.iter(qn("a:blip")):
                r_embed = blip.get(qn("r:embed"))
                if not r_embed:
                    continue
                try:
                    part = doc.part.related_parts[r_embed]
                    ext = _EXT_MAP.get(
                        part.content_type.split("/")[-1].lower(),
                        part.content_type.split("/")[-1].lower(),
                    )
                    parts.append(f"[IMG_{len(images)}]")
                    images.append((part.blob, ext))
                except (KeyError, AttributeError):
                    log.debug("DrawingML: не вдалося витягти rId=%s", r_embed)

            for imgdata in child.iter(f"{{{_VML_NS}}}imagedata"):
                r_id = imgdata.get(f"{{{_R_NS}}}id")
                if not r_id:
                    continue
                try:
                    part = doc.part.related_parts[r_id]
                    ext = _EXT_MAP.get(
                        part.content_type.split("/")[-1].lower(),
                        part.content_type.split("/")[-1].lower(),
                    )
                    parts.append(f"[IMG_{len(images)}]")
                    images.append((part.blob, ext))
                except (KeyError, AttributeError):
                    log.debug("OLE imagedata: не вдалося витягти rId=%s", r_id)

        para_texts.append("".join(parts))

    return "\n".join(para_texts), images

def _parse_section_header(text: str) -> tuple[str, str] | None:

    m = re.search(r'[«"](.*?)[»"]', text, re.DOTALL)
    if m:
        inner = m.group(1).strip()
    else:

        first_line = text.split("\n")[0].strip()

        code_m = re.match(r'^(\d+\.\d+(?:\.\d+)?\.?)\s+(.*)', first_line)
        if code_m:
            code = code_m.group(1).rstrip(".")
            name = code_m.group(2).strip()
            suffix = re.search(r'\(([a-zA-Zа-яА-ЯіІїЇєЄ])\)\s*$', name)
            if suffix:
                code = f"{code}({_normalize_suffix(suffix.group(1))})"
            return code, name
        return None

    parts = inner.split(None, 1)
    if not parts:
        return None
    code = parts[0].strip().rstrip(".")
    name = parts[1].strip() if len(parts) > 1 else code

    if not re.search(r'\([a-zA-Z]\)$', code):
        suffix = re.search(r'\(([a-zA-Zа-яА-ЯіІїЇєЄ])\)\s*$', name)
        if suffix:
            code = f"{code}({_normalize_suffix(suffix.group(1))})"
    return code, name

_CYR_SUFFIX_MAP = {"а": "a", "б": "b", "в": "c", "г": "d",
                   "с": "c", "е": "e", "d": "d"}

def _normalize_suffix(ch: str) -> str:
    ch = ch.lower()
    return _CYR_SUFFIX_MAP.get(ch, ch)

def _parse_question_row(cells: list[str]) -> dict | None:
    if len(cells) < 3:
        return None

    num_str = cells[0].strip()
    if num_str.isdigit():
        source_number = int(num_str)
    elif num_str in ("", "*", "–", "-"):
        source_number = 0
    else:
        return None

    if len(cells) > 1 and (cells[1].strip().isdigit() or cells[1].strip() == ""):
        text_col, ans_col = 2, 3
    else:
        text_col, ans_col = 1, 2

    if text_col >= len(cells) or ans_col >= len(cells):
        return None

    raw_text   = cells[text_col].strip()
    raw_answer = cells[ans_col].strip()

    if not raw_text or not raw_answer:
        return None

    _OPT_INLINE = re.compile(
        r'(?<=[^(\n])([ \t]*)([АВСABCавсБб]\s*[)\.][ \t])',
        re.UNICODE
    )
    raw_text = _OPT_INLINE.sub(r'\n\2', raw_text)
    lines = [ln.strip() for ln in raw_text.split("\n") if ln.strip()]
    question_text = ""
    options: dict[str, str] = {}

    for line in lines:

        m = re.match(r'^([АВСABCавсБб])\s*[)\.]\s*(.*)', line)
        if m:
            letter = m.group(1).upper()
            opt_text = m.group(2).strip()
            key = OPTION_KEYS.get(letter)
            if key:
                options[key] = opt_text
        else:
            if not question_text:
                question_text = line
            else:
                question_text += " " + line

    if not options and len(lines) >= 4:
        question_text = lines[0]
        option_lines = [ln for ln in lines[1:] if ln]
        if len(option_lines) >= 3:
            options["option_a"] = option_lines[0]
            options["option_b"] = option_lines[1]
            options["option_c"] = option_lines[2]

    if not all(k in options for k in ("option_a", "option_b", "option_c")):
        log.warning("Неповні варіанти відповідей для питання %d: %s", source_number, raw_text[:60])
        for k in ("option_a", "option_b", "option_c"):
            options.setdefault(k, "")

    correct = ANSWER_MAP.get(raw_answer.upper().strip(), raw_answer.upper().strip())
    if correct not in ("A", "B", "C"):
        log.warning("Незрозуміла правильна відповідь '%s' для питання %d", raw_answer, source_number)
        return None

    return {
        "source_number": source_number,
        "question_text": question_text,
        "option_a": options["option_a"],
        "option_b": options["option_b"],
        "option_c": options["option_c"],
        "correct_answer": correct,
    }

def _extract_category_from_heading(text: str) -> str | None:
    m = re.search(r'[Кк]атегор[іi]я\s*[:\-]?\s*([A-ZВ][0-9\.]+)', text)
    if m:
        raw = m.group(1).strip()

        raw = raw.replace("В", "B")

        raw = raw.rstrip(".")

        if raw.startswith("B1") or raw == "B2":
            raw = "T" + raw
        return raw
    return None

def parse_module_file(
    filepath: str,
    module_code: str,
    module_name: str,
    db_path: str = db.DB_PATH,
) -> dict[str, int]:
    log.info("Обробка файлу: %s", os.path.basename(filepath))

    doc = Document(filepath)

    module_id = db.insert_module(module_code, module_name, db_path)

    category_order: list[str] = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            cat = _extract_category_from_heading(para.text)
            if cat:
                category_order.append(cat)

    if not category_order:

        category_order = ["B1.1", "B1.3", "B2"]
        log.warning("Заголовки категорій не знайдені, використовую стандартні: %s", category_order)

    tables = doc.tables
    if len(tables) != len(category_order):
        log.warning(
            "Кількість таблиць (%d) не збігається з кількістю категорій (%d)",
            len(tables), len(category_order)
        )

    stats: dict[str, int] = {}

    for table_idx, table in enumerate(tables):
        if table_idx >= len(category_order):
            break

        cat_code = category_order[table_idx]

        if cat_code in ("TB1", "B1"):
            target_cats = ["B1.1", "B1.3"]
        else:
            target_cats = [cat_code]

        target_cat_ids = [db.insert_category(c, db_path) for c in target_cats]

        current_section_id: int | None = None
        saved_count = 0
        skipped_count = 0
        auto_number = 0

        for row_idx, row in enumerate(table.rows):
            cells      = [cell.text.strip() for cell in row.cells]
            raw_cells  = list(row.cells)

            if "кількість" in cells[0].lower() or (len(cells) > 2 and "кількість" in cells[2].lower()):
                header_text = cells[0] if "кількість" in cells[0].lower() else cells[2]
                parsed = _parse_section_header(header_text)
                if parsed:
                    sec_code, sec_name = parsed
                    current_section_id = db.insert_section(module_id, sec_code, sec_name, db_path)
                    log.debug("  Розділ: %s — %s", sec_code, sec_name)
                continue

            if cells[0].lower() in ("№ питання", "№"):
                continue
            if len(cells) > 1 and "питання та варіанти" in cells[1].lower():
                continue
            if len(cells) > 2 and "питання та варіанти" in cells[2].lower():
                continue

            if _is_excluded_row(row):
                skipped_count += 1
                continue

            if current_section_id is None:
                log.warning("Питання поза секцією у рядку %d, пропускаємо", row_idx)
                continue

            text_col = 2
            if len(cells) > 1 and not (cells[1].strip().isdigit() or cells[1].strip() == ""):
                text_col = 1

            if text_col < len(raw_cells):
                marked_text, cell_images = _extract_cell_content(raw_cells[text_col], doc)
                marked_cells = list(cells)
                marked_cells[text_col] = marked_text
            else:
                marked_cells = cells
                cell_images = []

            q = _parse_question_row(marked_cells)
            if q is None:
                continue

            if q["source_number"] == 0:
                auto_number += 1
                q["source_number"] = auto_number
            else:
                auto_number = q["source_number"]

            for cat_id in target_cat_ids:
                question_id = db.insert_question(
                    section_id=current_section_id,
                    category_id=cat_id,
                    source_number=q["source_number"],
                    question_text=q["question_text"],
                    option_a=q["option_a"],
                    option_b=q["option_b"],
                    option_c=q["option_c"],
                    correct_answer=q["correct_answer"],
                    db_path=db_path,
                )

                if cat_id == target_cat_ids[0]:
                    for context, field in [
                        ("question", q["question_text"]),
                        ("option_a", q["option_a"]),
                        ("option_b", q["option_b"]),
                        ("option_c", q["option_c"]),
                    ]:
                        local_idx = 0
                        for m_img in re.finditer(r'\[IMG_(\d+)\]', field):
                            global_idx = int(m_img.group(1))
                            if global_idx < len(cell_images):
                                img_data, ext = cell_images[global_idx]
                                db.insert_question_image(
                                    question_id, context, local_idx, ext, img_data, db_path
                                )
                                local_idx += 1

            saved_count += 1

        display_code = "/".join(target_cats)
        log.info("  Категорія %s: збережено %d питань (пропущено %d)", display_code, saved_count, skipped_count)
        for tc in target_cats:
            stats[tc] = saved_count

    return stats

def parse_all_modules(source_dir: str, db_path: str = db.DB_PATH) -> None:
    docx_files = sorted([
        f for f in os.listdir(source_dir)
        if f.endswith(".docx") and not f.startswith("~")
    ])

    if not docx_files:
        log.error("Не знайдено .docx файлів у %s", source_dir)
        return

    log.info("Знайдено %d файлів для обробки", len(docx_files))

    for filename in docx_files:
        filepath = os.path.join(source_dir, filename)

        m = re.search(r'[Мм]одуль\s*[№#]?\s*(\d+)', filename)
        if m:
            module_code = m.group(1)
        else:

            nums = re.findall(r'\d+', filename)
            module_code = nums[0] if nums else "0"

        name_m = re.search(r'[-–]\s*([А-ЯҐЄІЇа-яґєії ]+)', filename)
        module_name = name_m.group(1).strip() if name_m else f"Модуль {module_code}"

        parse_module_file(filepath, module_code, module_name, db_path)

    stats = db.get_db_stats(db_path)
    log.info("Імпорт завершено. Статистика БД: %s", stats)

if __name__ == "__main__":
    import sys

    db.init_db()

    if len(sys.argv) >= 3:

        fpath = sys.argv[1]
        mcode = sys.argv[2]
        mname = sys.argv[3] if len(sys.argv) > 3 else f"Модуль {mcode}"
        parse_module_file(fpath, mcode, mname)
    elif len(sys.argv) == 2:

        parse_all_modules(sys.argv[1])
    else:

        source = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "source_docs")
        parse_all_modules(source)

    print(db.get_db_stats())
